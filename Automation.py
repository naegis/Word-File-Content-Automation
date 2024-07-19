import os
import openai
import re
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#OPEN AI API
api_key = 'INSERT KEY HERE'
openai.api_key = api_key

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

'''
    THIS IS THE ONLY PART YOU WILL EDIT. TO GET THE TRAINING COURSE SUB TITLES, 
    YOU MUST USE CHATGPT, AND COPY ALL THE ITEMS UNDER "LIST OF [COURSE] TRAINING COURSES IN [COUNTRY]" AND SAY 
    "CAN YOU GET ONLY THE TRAINING COURSE TITLES ex: WordPress Training Course in Sinagpore, AND ENCLOSE THEM IN QUOTATION MARKS AND ADD A COMMA , AFTER EACH QUOTATION MARK?"
    THEN PASTE THE SUBCOURSE TITLES AND PUT THEM INSIDE THE training_courses = [here] list.
'''
#Sample Prompt
country = "Singapore"
maincourse = "Interpersonal Communication Training Courses in Singapore"
training_courses = [ 
"Conflict Management in Teams Training Courses in Singapore"]

#------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
#PROMPTING
def generate_content(prompt):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            n=1,
            stop=None,
            temperature=0.7,
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        print(f"An error occurred: {e}")
        return None


#WORD MAKING
def paste_content():
    doc = Document()

    #Main Course Title
    mainTitle = doc.add_paragraph(maincourse)
    run = mainTitle.runs[0]
    run.font.size = Pt(36)
    run.font.name = 'Arial'
    run.bold = True
    
    #function to format heading text
    def heading(course):
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(course)
        subtitle_run.bold = True
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(20)
        subtitle_run.font.color.rgb = RGBColor(21,95,129)  # Blue color


    # Split the prompt based on bullet points detection
    def line_splitter(prompt):
        bullet_point = '- '  # Define the bullet point format
        parts = prompt.split(bullet_point)  # Split by the bullet point format

        description = parts[0].strip()  # First part is the description
        bullet_list = [part.strip() for part in parts[1:]]  # Remaining parts are bullet points

        desc_paragraph = doc.add_paragraph()
        desc_run = desc_paragraph.add_run(description)
        desc_run.font.name = 'Arial'

        #Add the bullet points as a bulleted list
        for item in bullet_list:
            doc.add_paragraph(item, style='List Bullet')


    def create_formatted_document(prompt):
        content = generate_content(prompt)
        
        if content is None:
            print(f"Failed to generate content for {course}.")
            return
        
        # Split the content into sections
        lines = content.split('\n')
        description = []
        numbered_sections = []
        collecting_numbered_sections = False

        for line in lines:
            if re.match(r'^\d+\.', line.strip()):  # Check if the line starts with a number followed by a dot
                collecting_numbered_sections = True
                numbered_sections.append(line.strip())
            elif collecting_numbered_sections:
                numbered_sections.append(line.strip())
            else:
                description.append(line.strip())

        # Combine the description lines
        description_text = ' '.join(description).strip()
        
        # Add the description to the document
        desc_paragraph = doc.add_paragraph()
        desc_run = desc_paragraph.add_run(description_text)
        desc_run.font.name = 'Arial'

        # Add the numbered sections with bullet points
        for section in numbered_sections:
            if re.match(r'^\d+\.', section):  # Numbered section
                doc.add_paragraph(section, style='List Number')
            else:  # Bullet points under each numbered section
                doc.add_paragraph(section, style='List Bullet')



    for course in training_courses:
        #PROMPTS
        prompts = [
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Introduction Prompt: Please create a unique introduction for this title {course}. Lastly, please reiterate the title in the last sentence of the last paragraph. Make it 4-5 paragraphs.",
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Who Should Attend this {course} Prompt:Please create a 3 paragraph introduction For the training course title {course}. Lastly, please reiterate the title in the last sentence of the last paragraph. Can you also list down the titles/positions of people who might be interested about this, Just up their titles/positions no need for a description after the paragraphs.  do not add any other paragraph or description after the list  (ADD DASH BEFORE THE FIRST LIST ITEM NAME, ADD NEW LINE AFTER EVERY LIST ITEM AFTER), ",
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Course Duration Prompt: For the training course {course} title, create a 3 sentence paragraph introduction explaining the duration of the course training. Base the durations to the following: 3 full days, 1 day, half day, 90 minutes and 60 minutes. Please mention the title again in any of the 3 sentences.",
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Course Benefits Prompt: For the training course {course}, please create a 1 sentence description introducing the potential benefits of the training course. Afterwards, please list down 10 benefits of the {course} in bullet form. (ADD DASH BEFORE THE FIRST LIST ITEM NAME, ADD NEW LINE AFTER EVERY LIST ITEM AFTER) Also, do not add the word: \"Description:\" but still give me a description just without the word",
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Course Objectives Prompt For the training course {course}, Please create a 2 sentences description about the objectives of of the training course. Please mention the title again in any of the 2 sentences. Afterwards, please list down 12 objectives relating to the following benefits of the {course}. Make it different from the objectives in bullet form, (ADD DASH AT FIRST LIST ITEM, ADD NEW LINE AFTER EVERY LIST ITEM AFTER), AGAIN IT SHOULD BE 12 OBJECTIVES",
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Course Fees Prompt Please create a 3 sentence paragraph description for the possible course fees of {course}. Please put the course title somewhere in the sentences. Specify that there will be 4 pricing options but do not specifically specify any. (insert these bullets after the introduction) USD 679.97 For a 60-minute Lunch Talk Session. USD 289.97 For a Half Day Course Per Participant. USD 439.97 For a 1 Day Course Per Participant. USD 589.97 For a 2 Day Course Per Participant. Discounts available for more than 2 participants. ",
            f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Upcoming and Brochure Download Prompt: Please create a 3 sentence paragraph description for the possible upcoming updates or to avail brochures about the training course {course}. Please put the course title somewhere in the sentences. "
        ]

        #generate content and store into respective variables
        intro = generate_content(prompts[0])
        attendees = generate_content(prompts[1])
        duration = generate_content(prompts[2])
        benefits = generate_content(prompts[3])
        objectives = generate_content(prompts[4])
        fees = generate_content(prompts[5])
        brochure = generate_content(prompts[6])


        contentPrompt = f"Please only use British English. Be sure to use emotions. Be sure to connect with the reader. Course Content Prompt Please create a 2 sentence description for the possible course content of {course}. Please put the course title somewhere in the sentences. Take note of these when crafting the description {objectives}. Afterwards, create 12 sections using the same 12 objectives below, each with three topics composing of 1-2 sentences, pertaining to the following objectives of the course. Do not include colons, just the bullets. {objectives} Afterwards, create 12 sections (numbered list) using the same 12 objectives above, each with three topics (in bullet form) (Use \"-\" before their names as bullets) composing of 1-2 sentences, pertaining to the following objectives of the course. Do not include colons, just the bullets. Each of the 12 sections should have three topics composing of 1-2 sentences. And again, there should be a 2 sentence description for the possible course content of {course} at the start"
        content = generate_content(contentPrompt)

        if intro is None or attendees is None or duration is None or benefits is None or objectives is None or content is None or fees is None or brochure is None:
            print(f"Failed to generate content for {course}. Check error messages above.")
            continue



        

        #SubCourse Title
        title = doc.add_heading(f'{course}', level=1)
        title_format = title.runs[0].font
        title_format.size = Pt(28)

        
        #First main introductory bold paragraph (Cities where the course is also available)
        intro_paragraph = doc.add_paragraph()
        intro_paragraph.paragraph_format.line_spacing = 1.5
        intro_run = intro_paragraph.add_run(f"Our training course \"{maincourse}\" is also available in Orchard, Marina Bay, Bugis, Tanjong Pagar, Raffles Place, Sentosa, Jurong East, Tampines, Changi, and Woodlands.")
        intro_run.bold = True
        intro_run.italic = True
        intro_run.font.name = 'Arial'


        #Second Introduction
        doc.add_paragraph(intro)
        doc.paragraphs[-1].style = doc.styles['Normal']


        #Who Should Attend this [Sub course Title] Training Course in [Country] (Make this a subtitle but the format should be Arial, Font size: 20, Font color = blue)
        heading(f'Who Should Attend this {course}:')
        line_splitter(attendees)


        #Course Duration for [Sub course Title] Training Course in [Country] (Make this a subtitle but the format should be Arial, Font size: 20, Font color = blue)
        heading(f'Course Duration for {course}')
        duration_lines = duration.split('\n\n')
        intro_type = duration_lines[0]
        #Add intro paragraph
        doc.add_paragraph(intro_type)
        doc.add_paragraph("2 Full Days", style='List Bullet')
        doc.add_paragraph("9 a.m. to 5 p.m.", style='List Bullet')


        #Course Benefits of [Sub course Title] Training Course in [Country]  
        heading(f'Course Benefits of {course}:')
        line_splitter(benefits)
        

        #Course Objectives of [Sub course Title] Training Course in [Country]  
        heading(f'Course Objectives of {course}:')
        line_splitter(objectives)


        #Course Content for [Sub course Title] Training Course in [Country]  
        heading(f'Course Content for {course}')
        line_splitter(content)


        #Course Fees for [Sub course Title] Training Course in [Country]  
        heading(f'Course Fees for {course}:')
        fees_lines = fees.split('\n\n')
        fees_intro = fees_lines[0]
        #Add intro paragraph
        doc.add_paragraph(fees_intro)
        doc.add_paragraph("SGD 889.97 For a 60-minute Lunch Talk Session. ", style='List Bullet')
        doc.add_paragraph("SGD 389.97 For a Half Day Course Per Participant. ", style='List Bullet')
        doc.add_paragraph("SGD 589.97 For a 1 Day Course Per Participant. ", style='List Bullet')
        doc.add_paragraph("SGD 789.97 For a 2 Day Course Per Participant. ", style='List Bullet')
        discount_paragraph = doc.add_paragraph()
        discount_run = discount_paragraph.add_run("Discounts available for more than 2 participants.")
        discount_run.bold = True
        discount_paragraph.style = 'List Bullet'


        #Upcoming Course and Course Brochure Download for [Sub course Title] Training Course in [Country]  
        heading(f"Upcoming Course and Course Brochure Download for {course}")
        brochure_lines = brochure.split('\n\n')
        brochure_intro = brochure_lines[0]
        #Add intro paragraph
        doc.add_paragraph(brochure_intro)


    
    #Save the document, in my case, I save it directly into the downloads folder.
    downloads_path = os.path.join(os.path.expanduser("~"), "Downloads", f"{maincourse}.docx")
    try:
        doc.save(downloads_path)
        print(f"Document created successfully and saved to {downloads_path}")
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")



paste_content()
